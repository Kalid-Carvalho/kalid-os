from docx import Document
from docx.shared import Pt, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

COURIER = 'Courier New'
BODY_PT = Pt(10)
SP = Pt(3)

def fr(run, size=BODY_PT, bold=False, italic=False, color=None, underline=False):
    run.font.name = COURIER
    run.font.size = size
    run.font.bold = bold
    run.font.italic = italic
    run.font.underline = underline
    if color:
        run.font.color.rgb = color
    return run

def fp(para, align=WD_ALIGN_PARAGRAPH.JUSTIFY, sb=Pt(0), sa=SP, ls=1.15):
    fmt = para.paragraph_format
    fmt.alignment = align
    fmt.space_before = sb
    fmt.space_after = sa
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    fmt.line_spacing = ls
    return para

def ar(para, text, bold=False, italic=False, size=BODY_PT, color=None, underline=False):
    run = para.add_run(text)
    fr(run, size=size, bold=bold, italic=italic, color=color, underline=underline)
    return run

def rm_borders(tbl):
    tblPr = tbl._tbl.tblPr
    ex = tblPr.find(qn('w:tblBorders'))
    if ex is not None:
        tblPr.remove(ex)
    bdr = OxmlElement('w:tblBorders')
    for n in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        b = OxmlElement(f'w:{n}')
        b.set(qn('w:val'), 'none')
        b.set(qn('w:sz'), '0')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), 'auto')
        bdr.append(b)
    tblPr.append(bdr)

def no_cell_bdr(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    ex = tcPr.find(qn('w:tcBorders'))
    if ex is not None:
        tcPr.remove(ex)
    bdr = OxmlElement('w:tcBorders')
    for n in ['top', 'left', 'bottom', 'right']:
        b = OxmlElement(f'w:{n}')
        b.set(qn('w:val'), 'none')
        b.set(qn('w:sz'), '0')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), 'auto')
        bdr.append(b)
    tcPr.append(bdr)

def box_cell_bdr(cell, color='000000', sz=12):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    ex = tcPr.find(qn('w:tcBorders'))
    if ex is not None:
        tcPr.remove(ex)
    bdr = OxmlElement('w:tcBorders')
    for n in ['top', 'left', 'bottom', 'right']:
        b = OxmlElement(f'w:{n}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), str(sz))
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), color)
        bdr.append(b)
    tcPr.append(bdr)

def para_shd(para, fill):
    pPr = para._p.get_or_add_pPr()
    ex = pPr.find(qn('w:shd'))
    if ex is not None:
        pPr.remove(ex)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    pPr.append(shd)

# =====================
# Document setup
# =====================
doc = Document()

for p in list(doc.paragraphs):
    p._element.getparent().remove(p._element)

sec = doc.sections[0]
sec.page_width = Mm(210)
sec.page_height = Mm(297)
sec.top_margin = Mm(11)
sec.left_margin = Mm(12)
sec.right_margin = Mm(12)
sec.bottom_margin = Mm(45)
sec.footer_distance = Mm(5)

# =====================
# Footer
# =====================
footer = sec.footer
footer.is_linked_to_previous = False

for p in list(footer.paragraphs):
    p._element.getparent().remove(p._element)

blue_p = footer.add_paragraph()
fp(blue_p, align=WD_ALIGN_PARAGRAPH.CENTER, sb=Pt(3), sa=Pt(3), ls=1.4)
para_shd(blue_p, '1A4080')
ar(blue_p, 'Instituto Municipal de Planejamento Urbano – IMPLURB\n', italic=True, size=Pt(8.5), color=RGBColor(0xFF, 0xFF, 0xFF))
ar(blue_p, 'Av. Brasil, nº 2971 - Compensa (anexo da Prefeitura de Manaus) - CEP 69035-110 - Manaus - Amazonas', italic=True, size=Pt(8.5), color=RGBColor(0xFF, 0xFF, 0xFF))

gap = footer.add_paragraph()
fp(gap, sb=Pt(0), sa=Pt(4), ls=1.0)

gray_p = footer.add_paragraph()
fp(gray_p, align=WD_ALIGN_PARAGRAPH.LEFT, sb=Pt(2), sa=Pt(2), ls=1.3)
para_shd(gray_p, 'D4D4D4')
note = (
    'Este documento foi assinado digitalmente. Qualquer alteração invalidará o documento.\n'
    'Esta(s) assinatura(s) pode(m) ser verificada(s) em: https://fiscal.manaus.am.gov.br/fiscal/servlet/hwpvalidardocassinado (utilize a(s) chave(s) abaixo)\n'
    '1 - RADJA PEREIRA MAR. Em: Qui 10/06/2021 13:27:38. Chave para validação: B14M9P.894KDK.W984CG.FGGVK6.FW24A7.FX55B9\n'
    '2 - ISLANE RODRIGUES DE SOUZA. Em: Qui 10/06/2021 13:49:55. Chave para validação: EABMJZ.5RWZ5P.4D5JMT.U7XNI7.8JH18Z.FX55B9'
)
ar(gray_p, note, italic=True, size=Pt(6.5))

# =====================
# Helpers
# =====================
logo_path = 'c:/Users/kalid/Downloads/KalidCarvalho-os/dados/Eleicoes-2020-Veja-a-proposta-dos-candidatos-a-prefeitura-de-Manaus-1024x362.jpg'

def add_hdr(doc, ref_text):
    pr = doc.add_paragraph()
    fp(pr, align=WD_ALIGN_PARAGRAPH.RIGHT, sb=Pt(0), sa=Pt(3))
    ar(pr, ref_text, bold=True, size=Pt(10))

    ht = doc.add_table(rows=1, cols=2)
    rm_borders(ht)
    lc = ht.rows[0].cells[0]
    cc = ht.rows[0].cells[1]
    no_cell_bdr(lc)
    no_cell_bdr(cc)

    lp = lc.paragraphs[0]
    fp(lp, align=WD_ALIGN_PARAGRAPH.LEFT, sb=Pt(0), sa=Pt(0))
    if os.path.exists(logo_path):
        r = lp.add_run()
        r.add_picture(logo_path, height=Mm(26))

    inner = cc.add_table(rows=1, cols=3)
    rm_borders(inner)
    c0, c1, c2 = inner.rows[0].cells
    no_cell_bdr(c0)
    no_cell_bdr(c1)
    box_cell_bdr(c2)

    p0 = c0.paragraphs[0]
    fp(p0, align=WD_ALIGN_PARAGRAPH.LEFT, sb=Pt(0), sa=Pt(0))
    ar(p0, 'CERTIDÃO', bold=True, size=Pt(20))

    p1 = c1.paragraphs[0]
    fp(p1, align=WD_ALIGN_PARAGRAPH.LEFT, sb=Pt(5), sa=Pt(0))
    ar(p1, 'N°', bold=True, size=Pt(11))

    p2 = c2.paragraphs[0]
    fp(p2, align=WD_ALIGN_PARAGRAPH.CENTER, sb=Pt(3), sa=Pt(3))
    ar(p2, '4789/2021', bold=True, size=Pt(14), color=RGBColor(0xFF, 0, 0))

    sep = doc.add_paragraph()
    fp(sep, sb=Pt(2), sa=Pt(4))
    pPr = sep._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single')
    bot.set(qn('w:sz'), '6')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), '000000')
    pBdr.append(bot)
    pPr.append(pBdr)

def add_sig(doc):
    t = doc.add_table(rows=4, cols=1)
    rm_borders(t)

    r0 = t.rows[0].cells[0]
    no_cell_bdr(r0)
    p0 = r0.paragraphs[0]
    fp(p0, sb=Pt(8), sa=Pt(0), ls=1.15)
    ar(p0, 'E PARA CONSTAR EU, ')
    ar(p0, 'ISLANE RODRIGUES DE SOUZA, Gerente - GIT', underline=True)

    r1 = t.rows[1].cells[0]
    no_cell_bdr(r1)
    p1 = r1.paragraphs[0]
    fp(p1, sb=Pt(0), sa=Pt(0), ls=1.15)
    ar(p1, 'DO INSTITUTO MUNICIPAL DE PLANEJAMENTO URBANO - IMPLURB, LAVREI PRESENTE CERTIDÃO, QUE')

    r2 = t.rows[2].cells[0]
    no_cell_bdr(r2)
    p2 = r2.paragraphs[0]
    fp(p2, sb=Pt(0), sa=Pt(0), ls=1.15)
    ar(p2, 'VAI POR MIM ASSINADA E PELO(A) FUNCIONÁRIO(A) ')
    ar(p2, 'RADJA PEREIRA MAR', underline=True)

    r3 = t.rows[3].cells[0]
    no_cell_bdr(r3)
    p3 = r3.paragraphs[0]
    fp(p3, sb=Pt(0), sa=Pt(8), ls=1.15)
    ar(p3, 'AOS DIAS DO MÊS DE ')
    ar(p3, '10 de junho de 2021', underline=True)

# =====================
# PAGE 1
# =====================
add_hdr(doc, 'PAG CER: 4789/2021 PÁGINA 59/60')

p = doc.add_paragraph()
fp(p)
ar(p, 'TIPO DE CERTIDÃO:', bold=True)
ar(p, ' CERTIDÃO DE INFORMAÇÃO TÉCNICA PARA USO DO SOLO\n')
ar(p, 'EU,', bold=True)
ar(p, ' ISLANE RODRIGUES DE SOUZA, Gerente - GIT\n')
ar(p, 'CERTIFICO, TENDO EM VISTA AS ATRIBUIÇÕES QUE POR LEI A MIM SÃO CONFERIDAS E DE ACORDO COM O PROCESSO PROTOCOLADO N°', bold=True)
ar(p, ' 4789/2021 DE 25 de maio de 2021\n')
ar(p, 'REQUERIDO POR', bold=True)
ar(p, ' PETROMOR ALUGUEL DE SALAS E EQUIPAMENTOS LIMITADA\n')
ar(p, 'SITO À', bold=True)
ar(p, ' AVENIDA ANDRÉ ARAÚjO, nº 2721, ALEIXO, LOTEAMENTO PARQUE, MORADA DO SOL,\n')
ar(p, '         69060000,\n')
ar(p, 'DE PROPRIEDADE DE', bold=True)
ar(p, ' RODRIGO DOMENICO PETROSINO')

p = doc.add_paragraph()
fp(p)
ar(p, 'O(a) Interessado(a) solicita CERTIDÃO DE INFORMAÇÃO TÉCNICA PARA USO DO SOLO para o endereço acima descrito e conforme preceitua o Plano Diretor Urbano e Ambiental do Município de Manaus de 16/01/2014, informamos que:')

p = doc.add_paragraph()
fp(p)
ar(p, 'MATRÍCULA DO IPTU:', bold=True)
ar(p, ' 129069\n')
ar(p, 'ZONA:', bold=True)
ar(p, ' SETOR 11 (bairro Aleixo), FAIXA LINDEIRA DO CORREDOR URBANO ALEIXO, SEGMENTO ANDRÉ ARAÚjO.')

p = doc.add_paragraph()
fp(p, sb=Pt(5), sa=Pt(0))
ar(p, 'PARÂMETROS PERMITIDOS:', bold=True)

p = doc.add_paragraph()
fp(p)
ar(p, 'DIRETRIZES:', bold=True)
ar(p, ' Reforço do centro de comércio e serviços existente; integração de atividades comerciais e de serviços ao uso residencial.\n')
ar(p, 'USOS PERMITIDOS:', bold=True)
ar(p, ' Residencial unifamiliar e multifamiliar; comercial; de serviços; industrial.\n')
ar(p, 'ATIVIDADES PERMITIDAS:', bold=True)
ar(p, ' Tipo 1, Tipo 2, Tipo 3* e Tipo 4* (*exceto para o uso industrial).')

p = doc.add_paragraph()
fp(p, sb=Pt(5), sa=Pt(0))
ar(p, 'ATIVIDADE(S) SOLICITADA(S):', bold=True)

p = doc.add_paragraph()
fp(p, sb=Pt(0), sa=Pt(0))
ar(p, 'ATIVIDADE PRINCIPAL', bold=True)

p = doc.add_paragraph()
fp(p)
ar(p, 'CNAE/ATIVIDADE: 821130000 - Serviços combinados de escritório e apoio administrativo\n')
ar(p, 'USO/CLASSIFICAÇÃO: SERVIÇO TIPO 1')

p = doc.add_paragraph()
fp(p, sb=Pt(5), sa=Pt(0))
ar(p, 'ATIVIDADE(S) SECUNDÁRIA(S):', bold=True)

p = doc.add_paragraph()
fp(p)
ar(p, 'ATIVIDADE/CNAE:\n')
ar(p, '631190001 - Tratamento de dados, provedores de serviços de aplicação e serviços de hospedagem na internet\n')
ar(p, 'USO/CLASSIFICAÇÃO: SERVIÇO TIPO 1')

p = doc.add_paragraph()
fp(p)
ar(p, '829970701 - Salas de acesso à internet\n')
ar(p, '859960400 - Treinamento em desenvolvimento profissional e gerencial\n')
ar(p, '773310001 - Aluguel de máquinas e equipamentos para escritório\n')
ar(p, 'USO/CLASSIFICAÇÃO: SERVIÇO TIPO 2')

p = doc.add_paragraph()
fp(p)
ar(p, 'ANÁLISE: Permitido')

p = doc.add_paragraph()
fp(p)
ar(p, 'Observações gerais:\n')
ar(p, '1. O lote foi definido para USO RESIDENCIAL, estando o mesmo inserido no LOTEAMENTO PARQUE MORADA DO SOL, o qual obteve aprovação perante o Município em 24 de março de 1980, conforme demonstrado no base de dados da planta cadastral no ArqGIS;')

add_sig(doc)

# =====================
# PAGE 2
# =====================
doc.add_page_break()

add_hdr(doc, 'PAG CER: 4789/2021 PÁGINA 60/60')

p = doc.add_paragraph()
fp(p)
ar(p, '2. Nos termos da nova Lei Complementar de N° 012, de 17 de janeiro de 2019, § 4.° que diz: Nos casos de alterações de uso para as atividades Tipo 1 e 2, localizadas em loteamentos aprovados, necessitarão somente dos requisitos previstos no artigo 91 da LC 002, de 14 de janeiro de 2014, não sendo necessário aprovação prévia pelo CMDU e CTPCU, o pleito torna-se VIÁVEL;')

p = doc.add_paragraph()
fp(p)
ar(p, '3. De acordo com o Art. 96 § 4° da LC n° 014, de 17 de janeiro de 2019, por se tratar de um imóvel limítrofe a uma via caracterizada como CORREDOR URBANO, não cabendo a aplicação da cobrança da Outorga Onerosa de Alteração de Uso do Solo uma vez que trata de uma empresa de pequeno porte - EPP, ficando ISENTA do referido pagamento;')

p = doc.add_paragraph()
fp(p)
ar(p, '4. Deverá obedecer ao número de Vagas apresentadas para as atividades pleiteadas acima, sendo 05 (vagas) vagas para estacionamento devidamente demarcadas, sem obstrução do logradouro público, conforme disposto no Anexo IX - QUADRO DAS VAGAS DE GARAGEM E ESTACIONAMENTO, da Lei 2.402 de 16/01/2019;')

p = doc.add_paragraph()
fp(p)
ar(p, '5. Caso sejam constatados usos e/ou atividades diferentes destes autorizados e não permitidos para o local e/ou ampliações sem prévia aprovação, a presente certidão será automaticamente cancelada;')

p = doc.add_paragraph()
fp(p)
ar(p, '6. Deverão ser obedecidos os parâmetros das leis complementares n° 002, 003, 004, 005 e das leis n° 1.837, 1.838 e 1.839, de 16 de janeiro de 2014 do plano diretor urbano e ambiental do município de Manaus;')

p = doc.add_paragraph()
fp(p)
ar(p, '7. Esta certidão não garante o direito de construir, conforme disposto do parágrafo único, do artigo 14 da lei complementar n° 003/2014, não substitui o habite-se, bem como não reconhece direitos de propriedade;')

p = doc.add_paragraph()
fp(p)
ar(p, '8. Esta certidão não garante o direito de funcionamento, conforme artigo 9° da lei complementar n°005/2014, o funcionamento de qualquer estabelecimento comercial, industrial ou prestador de serviços, sem a necessária licença ou autorização, consiste em infração grave a referida lei, cabendo a fiscalização ao órgão licenciador das atividades econômicas do município;')

p = doc.add_paragraph()
fp(p)
ar(p, "9. Nos cursos d'água localizados na área urbana e de transição, será adotada faixa de proteção marginal de acordo com a legislação vigente;")

p = doc.add_paragraph()
fp(p)
ar(p, '10. Todo imóvel inserido, ainda que parcialmente, nos limites de área de preservação permanente APP, conforme lei federal n° 12.651/2012 de 25 de maio de 2012, deverá formalizar processo de avaliação junto ao órgão ambiental municipal ou estadual, para a possibilidade de obtenção de autorização específica;')

p = doc.add_paragraph()
fp(p)
ar(p, '11. Deverá ser mantido livre e desimpedido o passeio público, conforme disposto do inciso i, do artigo 83 da lei complementar n° 003/2014 e do artigo 38 da lei complementar n° 005/2014 e havendo geração de transtornos ao entorno, a presente certidão será cancelada;')

p = doc.add_paragraph()
fp(p)
ar(p, '12. Deverá atender ao parágrafo único do art. 82 da lei n° 1.838/2014, seção ii dos critérios e parâmetro para garagens e estacionamentos;')

p = doc.add_paragraph()
fp(p)
ar(p, '13. Deverá atender ao art.79 da lei n° 003/2014, subseção ii das garagens e estacionamentos para guarda de veículos;')

p = doc.add_paragraph()
fp(p)
ar(p, '14. Deverá atender ao Anexo IX - QUADRO DAS VAGAS DE GARAGEM E ESTACIONAMENTO, da Lei 2.402 de 16/01/2019.')

p = doc.add_paragraph()
fp(p)
ar(p, 'ESTAS INFORMAÇÕES NÃO PERDERÃO A VALIDADE, SALVO NO CASO DE ALTERAÇÃO SUPERVENIENTE DA LEGISLAÇÃO APLICÁVEL, CONFORME PARÁGRAFO ÚNCO DO ARTIGO 14 DA LEI COMPLEMENTAR N° 003 DE 16 DE JANEIRO DE 2014.')

p = doc.add_paragraph()
fp(p, align=WD_ALIGN_PARAGRAPH.CENTER, sb=Pt(5), sa=Pt(5))
ar(p, '*---------------------------------- FIM DESCRIÇÃO ----------------------------------*', bold=True)

add_sig(doc)

# =====================
# Save
# =====================
out = 'c:/Users/kalid/Downloads/KalidCarvalho-os/dados/certidao-output.docx'
doc.save(out)
print(f'Saved: {out}')
