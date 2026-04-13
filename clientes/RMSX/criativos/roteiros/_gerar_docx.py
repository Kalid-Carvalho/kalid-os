from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Estilos base
style_normal = doc.styles['Normal']
style_normal.font.name = 'Calibri'
style_normal.font.size = Pt(11)

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    if level == 1:
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(0, 0, 0)
    elif level == 2:
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(30, 30, 30)
    return p

def add_label(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(100, 100, 100)
    p.paragraph_format.space_after = Pt(2)
    return p

def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_script_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(40, 40, 40)
    return p

def add_divider(doc):
    p = doc.add_paragraph('─' * 60)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.runs[0]
    run.font.color.rgb = RGBColor(200, 200, 200)
    run.font.size = Pt(9)

# Capa
doc.add_heading('Roteiros — RMSX', 0)
p = doc.add_paragraph('Público: médicos que não têm tempo para marketing')
p.runs[0].font.color.rgb = RGBColor(80, 80, 80)
p = doc.add_paragraph('Plataforma: Meta Ads (Feed + Reels)')
p.runs[0].font.color.rgb = RGBColor(80, 80, 80)
doc.add_paragraph('')

# --- ROTEIRO 01 ---
add_divider(doc)
add_heading(doc, 'ROTEIRO 01 — Vídeo Curto de Venda Direta (15s)', level=1)
add_label(doc, 'Formato: #3 — Vídeo curto até 15s')
add_label(doc, 'Abertura: Dor direta')
add_label(doc, 'Onde gravar: Mesa de trabalho / escritório da RMSX')
add_label(doc, 'CTA no 7º segundo')
doc.add_paragraph('')

lines_01 = [
    "[0s] Você não virou médico pra gravar Reels.",
    "",
    "[3s] Mas seus concorrentes estão aparecendo todo dia.\n     E você, não.",
    "",
    "[7s] [CTA] A RSMX cuida de tudo enquanto você atende.\n     Clica aqui e fala com a gente.",
    "",
    "[10–15s] Sem contrato longo. Onboarding em 48h.",
]
for line in lines_01:
    add_script_block(doc, line)

add_label(doc, 'Instrução de gravação:')
add_body(doc, 'Câmera fixa, fundo limpo (escritório ou parede neutra). Fala direto, sem introdução. Começa já com "Você não virou médico..."')

# --- ROTEIRO 02 ---
add_divider(doc)
add_heading(doc, 'ROTEIRO 02 — Caixinha de Perguntas (melhor formato validado)', level=1)
add_label(doc, 'Formato: #1 — Caixinha de perguntas')
add_label(doc, 'Onde gravar: Mesa de trabalho, posição de autoridade')
doc.add_paragraph('')

add_heading(doc, 'Versão A — Pergunta sobre tempo', level=2)
lines_02a = [
    '[Texto na tela: "Mas eu não tenho tempo pra fazer marketing..."]',
    "",
    "[Vinicius responde na câmera]",
    "",
    "Exato. Você não deveria ter.\nIsso não é trabalho de médico. É trabalho de agência.",
    "",
    "A RSMX existe pra isso:\nvocê não grava, não posta, não configura anúncio,\nnão lê relatório.",
    "",
    "Você só atende.\nE no fim do mês a gente te diz quantos pacientes novos vieram do digital.",
    "",
    "Simples assim.\n[CTA] Clica aqui e começa essa semana.",
]
for line in lines_02a:
    add_script_block(doc, line)

doc.add_paragraph('')
add_heading(doc, 'Versão B — Pergunta sobre resultado', level=2)
lines_02b = [
    '[Texto na tela: "Já contratei agência e não funcionou..."]',
    "",
    "[Vinicius responde na câmera]",
    "",
    "Sim. Porque a maioria das agências trabalha pra qualquer negócio.\nA RSMX trabalha só pra médicos.",
    "",
    "A gente conhece o CFM, conhece o que pode e o que não pode,\ne sabe exatamente como trazer paciente novo pelo digital.",
    "",
    "Não é post bonito. É paciente na agenda.\n[CTA] Fala com a gente. Sem compromisso.",
]
for line in lines_02b:
    add_script_block(doc, line)

add_label(doc, 'Instrução de gravação:')
add_body(doc, 'Tela com a "pergunta" aparece primeiro (pode ser print de stories ou texto animado), depois Vinicius entra respondendo. Tom: parceiro, não vendedor.')

# --- ROTEIRO 03 ---
add_divider(doc)
add_heading(doc, 'ROTEIRO 03 — Vídeo de Autoridade (Dados)', level=1)
add_label(doc, 'Formato: #7 — Vídeo de autoridade')
add_label(doc, 'Onde gravar: Escritório ou mesa com computador/tela ao fundo')
add_label(doc, 'Duração: 30–45s')
doc.add_paragraph('')

lines_03 = [
    "[0s] Médico que trabalha com a RSMX devolve em média 14 horas por semana.",
    "",
    "[5s] 14 horas que você não vai gastar gravando vídeo de madrugada,\n     tentando entender o gerenciador de anúncios,\n     ou escrevendo legenda de post às 23h.",
    "",
    "[15s] No primeiro trimestre, nossos clientes tiveram\n      em média 72% mais pacientes novos vindos do digital.",
    "",
    "[25s] Não porque fizemos mágica.\n      Porque a gente faz isso todo dia, só pra médicos,\n      desde antes de você ouvir falar em tráfego pago.",
    "",
    "[35s] [CTA] Se você quer esses números no seu consultório,\n      clica aqui e agenda uma conversa.\n      Sem pitch. Sem enrolação. Só uma conversa direta.",
]
for line in lines_03:
    add_script_block(doc, line)

add_label(doc, 'Instrução de gravação:')
add_body(doc, 'Números aparecem na tela conforme você fala (edição simples). Fala pausada, com confiança. Não precisa de teleprompter.')

# --- ROTEIRO 04 ---
add_divider(doc)
add_heading(doc, 'ROTEIRO 04 — Vídeo Nichado (Ambiente Médico)', level=1)
add_label(doc, 'Formato: #8 — Vídeo nichado com ambiente do cliente')
add_label(doc, 'Onde gravar: Corredor de hospital, clínica, consultório ou sala de espera')
add_label(doc, 'Duração: 20–30s')
doc.add_paragraph('')

lines_04 = [
    "[0s — filmado no corredor de uma clínica ou consultório]",
    "",
    "Você construiu esse consultório com anos de estudo e dedicação.",
    "",
    "[5s] Mas se alguém der uma busca no Google agora,\n     o que vai aparecer?",
    "",
    "[10s] Um site desatualizado.\n      Dois posts de três meses atrás.\n      E o concorrente no topo do resultado.",
    "",
    "[18s] A RSMX resolve isso enquanto você continua aqui,\n      fazendo o que você sabe fazer.",
    "",
    "[25s] [CTA] Clica e fala com a gente hoje.",
]
for line in lines_04:
    add_script_block(doc, line)

add_label(doc, 'Instrução de gravação:')
add_body(doc, 'Câmera começa mostrando o ambiente (passando pelo corredor ou filmando a sala), depois Vinicius entra no quadro ou faz narração em off. O ambiente médico é o que chama a atenção no feed.')

# --- ROTEIRO 05 ---
add_divider(doc)
add_heading(doc, 'ROTEIRO 05 — Vídeo Institucional (Apresentação da RMSX)', level=1)
add_label(doc, 'Formato: #5 — Vídeo institucional')
add_label(doc, 'Onde gravar: Escritório da RMSX, com equipe se houver')
add_label(doc, 'Duração: 30–45s')
doc.add_paragraph('')

lines_05 = [
    "[0s] Eu sou Vinicius, fundador da RMSX.",
    "",
    "[3s] A gente existe porque eu via médicos incríveis,\n     especialistas de verdade,\n     que precisavam passar hora gravando Reels pra não ficar pra trás no digital.",
    "",
    "[10s] Isso não faz sentido.\n      Médico bom deveria estar atendendo paciente, não aprendendo a mexer no Meta.",
    "",
    "[15s] A RMSX cuida de toda a presença digital do médico:\n      site, redes sociais, anúncios no Google e no Instagram,\n      tudo dentro das normas do CFM.",
    "",
    "[25s] Você fala com a gente uma vez, 30 minutos.\n      Em 15 dias, tudo no ar.\n      Depois disso, você só atende.",
    "",
    "[35s] [CTA] Clica aqui se quiser saber como funciona.",
]
for line in lines_05:
    add_script_block(doc, line)

add_label(doc, 'Instrução de gravação:')
add_body(doc, 'Pode mostrar a equipe trabalhando ao fundo. Vinicius fala direto pra câmera, tom de parceiro. Cortes curtos no meio se quiser dinamismo.')

# --- ROTEIRO 06 ---
add_divider(doc)
add_heading(doc, 'ROTEIRO 06 — Vídeo Banco + Narração (fácil de produzir)', level=1)
add_label(doc, 'Formato: #9 — Vídeo de banco + narração')
add_label(doc, 'Produção: Canva ou CapCut. Sem câmera.')
add_label(doc, 'Vídeos de banco: consultório cheio, médico atendendo, celular com feed, pessoa digitando')
doc.add_paragraph('')

lines_06 = [
    "[Narração — voz do Vinicius ou narrador]",
    "",
    "Enquanto você tá atendendo,\nalguém está no Google procurando exatamente o que você faz.",
    "",
    "[Vídeo: pessoa no celular, rolando feed]",
    "",
    "A pergunta é: eles estão te achando?",
    "",
    "[Vídeo: consultório vazio vs. consultório cheio]",
    "",
    "Médicos que investem em presença digital\ntrazem em média 72% mais pacientes novos no primeiro trimestre.",
    "",
    "[Vídeo: pessoa satisfeita, agenda cheia]",
    "",
    "A RSMX faz isso acontecer.\nSem você precisar gravar um Reels sequer.",
    "",
    "[CTA]\nClica e fala com a gente. Onboarding em 48h.",
]
for line in lines_06:
    add_script_block(doc, line)

add_label(doc, 'Instrução de produção:')
add_body(doc, 'Montar no Canva com vídeos do Pexels (buscar: "doctor", "clinic", "medical consultation"). Narração por cima. Legendas no estilo Reels.')

# --- Prioridade ---
add_divider(doc)
add_heading(doc, 'Prioridade de Execução', level=1)

table = doc.add_table(rows=7, cols=3)
table.style = 'Table Grid'

headers = ['Prioridade', 'Roteiro', 'Motivo']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    cell.paragraphs[0].runs[0].bold = True

rows_data = [
    ('1', 'Roteiro 01 — Vídeo Curto 15s', 'Mais rápido de gravar. Testar abertura primeiro.'),
    ('2', 'Roteiro 06 — Banco + Narração', 'Sem câmera. Pronto em 1h. Bom pra testar copy.'),
    ('3', 'Roteiro 02A — Caixinha de Perguntas', 'Melhor formato validado. Requer um pouco mais de edição.'),
    ('4', 'Roteiro 03 — Autoridade + Dados', 'Exige os dados reais confirmados.'),
    ('5', 'Roteiro 04 — Nichado em Clínica', 'Depende de acesso ao ambiente médico.'),
    ('6', 'Roteiro 05 — Institucional', 'Produção maior. Bom para remarketing.'),
]

for i, (p, r, m) in enumerate(rows_data):
    row = table.rows[i + 1]
    row.cells[0].text = p
    row.cells[1].text = r
    row.cells[2].text = m

output_path = r'c:/Users/kalid/Downloads/KalidCarvalho-os/clientes/RMSX/criativos/roteiros/roteiros-vinicius.docx'
doc.save(output_path)
print(f"Arquivo salvo em: {output_path}")
